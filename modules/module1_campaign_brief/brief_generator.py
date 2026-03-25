import anthropic
import streamlit as st
from brand_voice import BRAND_VOICE_FINGERPRINT


def get_client():
    """Get Anthropic client using secrets.toml or environment variable."""
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        import os
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise ValueError("Anthropic API key not found. Add it to .streamlit/secrets.toml")
    return anthropic.Anthropic(api_key=api_key)


def generate_brief(campaign_data: dict) -> str:
    client = get_client()

    prompt = f"""
You are the WAI USA Marketing Intelligence system. Generate a complete, structured campaign brief based on the intake data below.

{BRAND_VOICE_FINGERPRINT}

## Campaign Intake Data:
- Campaign Name: {campaign_data.get('campaign_name', 'N/A')}
- Campaign Type: {campaign_data.get('campaign_type', 'N/A')}
- Primary Goal: {campaign_data.get('primary_goal', 'N/A')}
- Target Audience: {campaign_data.get('target_audience', 'N/A')}
- Key Message: {campaign_data.get('key_message', 'N/A')}
- Channels: {', '.join(campaign_data.get('channels', []))}
- Timeline: {campaign_data.get('timeline', 'N/A')}
- Budget/Resources: {campaign_data.get('budget', 'N/A')}
- Programs Involved: {', '.join(campaign_data.get('programs', [])) if campaign_data.get('programs') else 'None specified'}
- Additional Notes: {campaign_data.get('notes', 'None')}

## Generate a Campaign Brief with these exact sections:

**1. CAMPAIGN OVERVIEW**
One punchy paragraph that captures what this campaign is, why it matters, and what success looks like. WAI USA voice. No fluff.

**2. CAMPAIGN OBJECTIVE**
Primary objective (one clear sentence) + 2-3 supporting objectives as bullet points.

**3. TARGET AUDIENCE**
Who we're reaching, what they care about, and what motivates them to act.

**4. KEY MESSAGES**
3-4 core messages, each as a single punchy line. These should feel like copy hooks.

**5. CHANNEL STRATEGY**
For each selected channel, write 1-2 sentences on how we use it specifically for this campaign.

**6. CONTENT PILLARS**
3 content pillars for this campaign. Each: pillar name + 1-sentence description + 1 example content idea.

**7. CALLS TO ACTION**
2-3 specific CTAs for this campaign. Make them action-forward and channel-appropriate.

**8. SUCCESS METRICS**
4-5 specific KPIs to track. Be concrete (e.g., "LinkedIn post impressions" not just "engagement").

**9. TIMELINE & MILESTONES**
Break the campaign timeline into phases with key milestones. Be specific to the dates/duration provided.

**10. APPROVAL NOTES FOR LEADERSHIP**
2-3 bullet points flagging any decisions, approvals, or inputs needed from leadership before execution begins. Do not address any specific person by name.

Write in WAI USA voice throughout. Be specific, not generic. This brief should be ready for leadership review before execution.
"""

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=2500,
        messages=[{"role": "user", "content": prompt}]
    )

    return message.content[0].text


def generate_docx(campaign_data: dict, brief_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()
    section = doc.sections[0]
    section.page_width = 12240
    section.page_height = 15840
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    title = doc.add_heading(f"Campaign Brief: {campaign_data.get('campaign_name', 'Untitled')}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0x6B)
        run.font.size = Pt(20)

    sub = doc.add_paragraph(
        f"WAI USA Marketing OS  |  {campaign_data.get('campaign_type', 'N/A')}  |  {campaign_data.get('timeline', 'TBD')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for line in brief_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('**') and line.endswith('**') and any(c.isdigit() for c in line[:5]):
            h = doc.add_heading(line.strip('*').strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0D, 0x47, 0x6B)
        elif line.startswith('- ') or line.startswith('• '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(11)
        elif '**' in line:
            p = doc.add_paragraph()
            for i, part in enumerate(line.split('**')):
                if not part:
                    continue
                run = p.add_run(part)
                run.bold = (i % 2 == 1)
                run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        "Generated by WAI USA Marketing OS  |  Requires leadership approval before execution"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
